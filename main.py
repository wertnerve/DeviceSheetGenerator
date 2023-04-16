import openpyxl
import docx
import time
# library used for reading excel files
from openpyxl import load_workbook
from docx import Document
#phys por t- device a port


#device a port - device z

#phys port z > device z port

#creates the spaces between each section
def sectionHead():
    return '='*70
#spaces between values on the sheet
def spaceGen(length):
    return " " * int(length)

#removes None values from list
def removeNone(row):
    conv = lambda i: i or ''
    res = [conv(i) for i in row]
    return res

#takes in all values read from each device and generates a word doc with the device's information
def generateDeviceSheet(cabinet, device, hostname, psu, pporta, pmt, dZ, pportz, rackU, mount, depth, failLocal,
                        failRemote, n, cTitles):

    #convert None to a string ''
    print(f"The hostname is {hostname}")
    print(f"The device is {device}")

    n = removeNone(n)

    filename = str(hostname) + ".doc"
    if '?' in filename:
        print("Replacing unsuported value")
        filename = filename.replace('?','TBD')
    print("Filename is",filename)
    #filepath = "C:\Temp\Data Center Script\Device Sheets\ " + filename
    #filepath = "S:\TIG\DataCenter\DeviceSheets\ " + filename

    print("Creating new file:",filename)

    print("Here is the collected data from the  device we just parsed:")

    print("\n")

    spaceLen = 62
    maxSpace =110
    doc = Document()

    #make the document landscape
    for section in doc.sections:
        section.page_width = docx.shared.Inches(11)
        section.page_height = docx.shared.Inches(8.5)
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE

    rowLine = f"DEVICE A{spaceGen(maxSpace)}HOST NAME"
    rowline2 = f"New Cabinet{spaceGen(spaceLen-len('cabinet       '))}Mount F/B/S "

    print(sectionHead())
    print(rowLine)
    doc.add_paragraph(rowLine)

    #static variables
    print(f"{device} {spaceGen(30-len(device))} {hostname} \n")
    print("=" * len(rowLine))
    print(rowline2)
    print(f" {cabinet} {spaceGen(50)} {mount}\n")

    doc.add_paragraph(f"{device}{spaceGen(maxSpace-len(device))}{hostname}")
    doc.add_paragraph(f"{sectionHead()}")
    doc.add_paragraph(rowline2)
    if mount is not None:
        doc.add_paragraph(f"{cabinet} {spaceGen(spaceLen)} {mount}")
    else:
        doc.add_paragraph(f"{cabinet} {spaceGen(spaceLen - 1)} {mount}")
    doc.add_paragraph(f"{sectionHead()}")

    print("=" * len(rowLine))
    print(f"Physical Port A{spaceGen(spaceLen - 10)}Port/Media Type\n")

    doc.add_paragraph(f"Device A Port {spaceGen(spaceLen - 30)}Port/Media Type{spaceGen(spaceLen-30)}  Device Z    {spaceGen(spaceLen-30)}Device Z Port")

    listLen = len(pporta)
    count = 0
    spaceLen = 50
    while count < listLen:
        if dZ[count] is not None and pmt[count] is not None:
             portalen = 0
             if pporta[count] is not None:
                 portalen = len(pporta[count])
             else:
                 portalen = 4

             rowLine = str(f"{count+1}. {pporta[count]}{spaceGen(spaceLen - portalen)}"
                      f"{str(pmt[count])}{spaceGen(spaceLen - len(dZ[count]))}{str(deviceZ[count])}{spaceGen(spaceLen - len(dZ[count]))}{str(pportz[count])}")
        else:
             rowLine = str(f"{count+1}. {pporta[count]}{spaceGen(spaceLen - 4)}"
                          f"{str(pmt[count])}{spaceGen(spaceLen-4)}{str(deviceZ[count])}{spaceGen(spaceLen-4)}{str(pportz[count])}")
        print(rowLine)
        doc.add_paragraph(rowLine)
        count+=1

    print(sectionHead())
   # file.write(f"{sectionHead()} \n")

    print((f"Device A Ports{spaceGen(spaceLen-10)}Physical Port Z\n"))
   # file.write(f"Device A Ports{spaceGen(spaceLen-12)}Physical Port Z\n \n")


    print(sectionHead())
    doc.add_paragraph(sectionHead())
    doc.add_paragraph(f"Notes:")

    noNotes = True
    #account for if the note tab in the first row for the device is empty, skip over it
    skip = True
    if n[0] != '':
        skip = False
    index = 0
    for item in n:
        if skip:
            skip = False
        else:
            index += 1
        if item != '' and not skip:
           doc.add_paragraph(f"{index}. {item}")
           noNotes = False
       # else:
          #  doc.add_paragraph(f"{index}.{item}")
    if noNotes:
        doc.add_paragraph("No Notes")

    doc.save("S:\TIG\DataCenter\DeviceSheets\ " + filename)
    print("File created")

if __name__ == '__main__':
    print("Starting program...")
    #this matrix contains the hostname of each indivdual device and what model it is
    deviceList = []
    deviceType = []
    deviceCount = 0
    # Spreadsheet is stored in the below filepath on FE2694
    data_file = 'S:\TIG\DataCenter\Test-DC-Data2.xlsx'
    # load the spreadsheet into the program
    wb = load_workbook(data_file)
    #Todo READ THE BELOW NOTE
    #######NOTE!!!! MAKE SURE THE SHEETNAME IS CORRECT HERE, AND THAT THE COLUMN NAMES MATCH THE ONES IN THE FILE LISTED ABOVE
    ws = wb['Sheet1']
    all_rows = list(ws.rows)
    rowLen = len(all_rows)

    print(f"Found {rowLen} rows of data.")
    # this keeps track of which row were up to, and will be used to check the upcoming row for a duplicate hostname, which is where we break between devices
    rowCounter = 1
    # we want to start our intenal rowcounter where the device name sstart, which is row 2. Rmemeber, row 1 is where th3e column names are
    rowNum = 2
    # this pritns each value of the top row in the sheet, basically prints the titles
    columnTitles = []

    # below are the lists we use for our nonstatic variables,
    # AKA the different values for each device that need to be recorded more than onece
    # physicalPortA
    physicalPA = []
    # Port Media Type
    portMedia = []
    deviceZ = []
    # physcial Port Z
    physicalPZ = []
    notes = []
    # bolean vairbale to dtermine if weve collected the static variables, such as hostname, device, physical traits that only need to be recorded once
    staticCollected = False
    staticValues = []

  #inactive loop, gets the column names
    for column in all_rows[0]:
        columnTitles.append(column.value)
    print("\n")

    valueLists = [physicalPA, portMedia, deviceZ, physicalPZ, notes, staticValues]
    # keeps track of the upcoming row/hostname
    nextRow = all_rows[rowCounter]

    # this loop runs through the whole spreadsheet, starting right past the column titles on row 1
    for row in all_rows[1:rowLen]:
        if rowNum > 46:
            print("Current row:", rowNum)
            print("Static Values",staticValues)
            print("List values",valueLists)
        # For the first time we detect a new device, pick up its one-time/static values, dont refresh until we hit a new device
        if staticCollected is False:
            staticValues = [row[1].value, #New Cabinet
                            row[3].value, #Device A
                            row[4].value, #Hostname
                            row[5].value, #PSUCount
                            row[6].value, #RackUnits
                            row[12].value, #Mount F/B/S
                            row[8].value, #Depth
                            row[9].value, #Failover Local
                            row[10].value] #Failover Remote
            staticCollected = True

        rowCounter += 1
        rowNum += 1
        if rowNum > rowLen:
          break
        else:
            # if we havent detected a new device, update each list with the new info, and move to the next row in the cache
            try:
                nextRow = all_rows[rowCounter]
            except:
                print("Weve hit row", rowNum)
                break

            physicalPA.append(row[6].value)
            portMedia.append(row[7].value)
            deviceZ.append(row[8].value)
            physicalPZ.append(row[9].value)


            notes.append(row[17].value)
        deviceCount += 1

        # once the program detects a new hostname on the upcoming row, its time to export thsi devices information and start witht he next one
        #the second condition covers when were at the end of the file
        if nextRow[4].value is not row[4].value:
            print(f"The current hostname is {row[3].value} , which has appeared in the sheet {deviceCount} times")
            print(
                "The upcoming row, Row#",rowNum,", has a different Hostname! No further rows for current device. Generating results...")
            print("\n")
            deviceList.append(row[4].value)
            deviceType.append([row[3].value])
            # remove the 'None" that gets read at the start of each device for the lists, bypass if list lenght < 2
            if physicalPA[0] is None:
                physicalPA.pop(0)
                deviceZ.pop(0)
                portMedia.pop(0)
                physicalPZ.pop(0)
            # pass all values we've gathered to the sheet generator
            generateDeviceSheet(staticValues[0], staticValues[1], staticValues[2], staticValues[3], physicalPA,
                                portMedia, deviceZ, physicalPZ,
                                staticValues[4], staticValues[5], staticValues[6], staticValues[7], staticValues[8],
                                notes,columnTitles)
            if (rowNum) > rowLen:
                print("We've reached the end of the file! We are at row", rowNum , "... Terminating program")
            else:
                print(f"Thats all for this device! Moving onto HOSTNAME {nextRow[3].value} , starting at row {rowNum}")
            print("\n")
            deviceCount = 0
            # empty the lists of all the values for the device we just parsed through, both static and nonstatic
            for item in valueLists:
                item.clear()
            staticCollected = False


    dCount = 0
    filepath = "S:\TIG\DataCenter\DeviceSheets\Summary.doc"
    file = open(filepath, "w")
    print(f"In total, we collected information on {len(deviceList)} devices. Their HOSTNAMES and MODEL are listed below:")
    file.write(f"In total, the script parsed {len(deviceList)} devices. \nTheir HOSTNAMES and MODEL TYPES are listed below: \n{sectionHead()} \n \n")
    for hostname in deviceList:

        line = f"HOSTNAME: {hostname} \n"
        print(line)
        file.write(line)

        line = f"DEVICE MODEL: {deviceType[dCount]} \n"
        print(line)
        file.write(line)
        file.write("\n")
        file.write(f"{sectionHead()} \n \n")
        dCount+=1
    file.close()
# maybe a stack wouldve been more efficent here, but hey, what can you do hahaha
